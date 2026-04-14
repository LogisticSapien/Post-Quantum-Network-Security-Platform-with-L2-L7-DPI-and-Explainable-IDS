"""Convert Quantum Sniffer report markdown to styled DOCX (v3)."""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x0B, 0x2A, 0x5C)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(22)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(16)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(13)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


def add_styled_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0B2A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    ncols = len(headers)
    for r_idx, row in enumerate(rows):
        # Normalize row to have exactly ncols entries
        norm = list(row[:ncols]) + [''] * max(0, ncols - len(row))
        for c_idx in range(ncols):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(norm[c_idx]))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            bg = "F2F6FC" if r_idx % 2 == 0 else "FFFFFF"
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_bold_para(label, value):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Calibri'
    p.add_run(value).font.name = 'Calibri'


# ══════════════════════════════════════════════════════════════════
# READ THE MARKDOWN AND CONVERT
# ══════════════════════════════════════════════════════════════════

with open('Quantum_Sniffer_v2_Technical_Report.md', 'r', encoding='utf-8') as f:
    md = f.read()

# Replace v2.0 with v3.0
md = md.replace('v2.0', 'v3.0').replace('v2 ', 'v3 ').replace('Version:** 2.0', 'Version:** 3.0')

lines = md.split('\n')
i = 0
in_code = False
code_buf = []
in_table = False
table_headers = []
table_rows = []

def flush_table():
    global in_table, table_headers, table_rows
    if in_table and table_headers:
        add_styled_table(table_headers, table_rows)
    in_table = False
    table_headers = []
    table_rows = []

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip HTML-style anchors
    if stripped.startswith('<') and stripped.endswith('>'):
        i += 1
        continue

    # Code blocks
    if stripped.startswith('```'):
        if in_code:
            flush_table()
            add_code_block('\n'.join(code_buf))
            code_buf = []
            in_code = False
        else:
            flush_table()
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # Table rows
    if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        # Skip separator rows
        if all(set(c) <= set('-: ') for c in cells):
            i += 1
            continue
        if not in_table:
            in_table = True
            table_headers = cells
        else:
            table_rows.append(cells)
        i += 1
        continue
    else:
        flush_table()

    # Horizontal rule
    if stripped == '---' or stripped == '***':
        # skip
        i += 1
        continue

    # Headings
    if stripped.startswith('#'):
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = min(len(m.group(1)), 3)
            title = m.group(2).strip()
            # Clean markdown links
            title = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', title)
            title = title.replace('**', '')
            doc.add_heading(title, level=level)
            i += 1
            continue

    # Empty line
    if not stripped:
        i += 1
        continue

    # Bullet points
    if stripped.startswith('- ') or stripped.startswith('* '):
        text = stripped[2:]
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        # Handle bold markers
        p = doc.add_paragraph(style='List Bullet')
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            else:
                run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        i += 1
        continue

    # Numbered items
    m = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if m:
        text = m.group(2)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        p = doc.add_paragraph(style='List Number')
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    # Regular paragraph
    text = stripped
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    if text.startswith('**') and text.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(text[2:-2])
        run.bold = True
        run.font.name = 'Calibri'
    else:
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Calibri'
            else:
                # Handle inline code
                sub_parts = re.split(r'(`[^`]+`)', part)
                for sp in sub_parts:
                    if sp.startswith('`') and sp.endswith('`'):
                        run = p.add_run(sp[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x0B, 0x2A, 0x5C)
                    else:
                        run = p.add_run(sp)
                        run.font.name = 'Calibri'
    i += 1

flush_table()

# ── Save ──
output = 'Quantum_Sniffer_v3_Technical_Report.docx'
doc.save(output)
print(f"Saved: {output}")
